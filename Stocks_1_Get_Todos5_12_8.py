import requests
import pandas as pd
import os
from datetime import datetime
import time
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import smtplib
from email.message import EmailMessage
from openpyxl import load_workbook
from openpyxl.styles import Font
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter
from pathlib import Path

# Esta version extrae los datos y saca la 
# DIFERENCIA, MOVIMIENTOS Y STOCK CRITCO
#Solo agrega los datos cuando hay diferencia, en caso contrario, no guarda nada.
#Cambiamos delimitadores de CSV por ";".
#Agregamos columna DIAS STOCK ACTUAL, STOCK MAXIMO, etc. a hoja RESUMEN.
#Agregamos alerta de Stock Critico cada vez que un producto entra en él.



# ---------------Obtención datos de BSALE---------------

# 📌 Configuración del servidor BSALE
BSALE_URL = "https://api.bsale.io/v1/stocks.json"
TOKEN = "a1f4df6fb5b62913421416ef30f7b91bcd15d759"
LIMIT = 50

# 📌 Definir intervalo de actualización en segundos
INTERVALO_SEGUNDOS = 60  # Cambia esto según la frecuencia deseada

#def actualizar_stock():

# 📌 Encabezados para autenticación
headers = {
    "access_token": TOKEN,
    "Content-Type": "application/json"
}

# 📌 Lista donde guardaremos todos los productos
todos_los_productos = []

# 📌 Inicializar paginación - offset = n°de página (comienza en la pág. 0)
offset = 0

# Obtener la fecha y hora actual solo UNA VEZ antes de recorrer los productos
timestamp_actual = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

while True:
    response = requests.get(f"{BSALE_URL}?limit={LIMIT}&offset={offset}&expand=variants,offices,product", headers=headers)

    if response.status_code == 200:
        data = response.json()
        data1 = data.get("items", [])

        if not data1:
            print("\n🔹 No hay más productos para obtener. Finalizando descarga.")
            break  

        # 📌 Extraer solo la información necesaria
        for dato in data1:
            data2 = {
                "Fecha": timestamp_actual,  
                "Producto": dato.get("variant", {}).get("product", {}).get("name"),
                "SKU": dato.get("variant", {}).get("code"),
                "UND": dato.get("variant", {}).get("description"),
                "Bodega": dato.get("office", {}).get("name"),
                "Stock": dato.get("quantity")
            }
            todos_los_productos.append(data2)  

        offset += LIMIT  

    else:
        print(f"\n❌ Error {response.status_code}: {response.text}")
        break  

print(f"\n✅ Total de productos obtenidos: {len(todos_los_productos)}")

# 📌 Convertir lista a DataFrame
df_nuevo = pd.DataFrame(todos_los_productos)

print("Fecha es NA después de la extracción:")
print(df_nuevo[df_nuevo["Fecha"].isna()])





# -----------------------------RUTAS-------------------------------------------

# 📌 Rutas para guardar los archivos
BASE_DIR = Path(__file__).resolve().parent
ruta_carpeta = BASE_DIR / "Proyecto" / "Stock"
os.makedirs(ruta_carpeta, exist_ok=True)
ruta_csv = os.path.join(ruta_carpeta, "historial_stock.csv")
ruta_productos_unicos = os.path.join(ruta_carpeta, "productos_unicos.csv")
ruta_excel = os.path.join(ruta_carpeta, "historial_stock.xlsx")
ruta_word = os.path.join(ruta_carpeta,"Informe Stock Productos.docx")
ruta_word2 = os.path.join(ruta_carpeta,"Informe Stock Crítico.docx")
ruta_word3 = os.path.join(ruta_carpeta,"Informe Stock Crítico - MP.docx")
ruta_word4 = os.path.join(ruta_carpeta,"Informe Stock Crítico - PT.docx")
ruta_word_bv = os.path.join(ruta_carpeta,"Informe Stock Crítico - BV.docx")
ruta_word_compras = os.path.join(ruta_carpeta,"Informe Stock Crítico - COMPRAS.docx")
ruta_ventas = os.path.join(ruta_carpeta,"Ventas.xlsx")
ruta_consumos = os.path.join(ruta_carpeta,"Consumos.xlsx")
ruta_stock_critico = os.path.join(ruta_carpeta, "Stock_Critico.xlsx")  # 📌 Ruta del archivo de stock crítico
ruta_stock_critico_compilado = os.path.join(ruta_carpeta, "STOCK CRÍTICO COMPILADO.xlsx")  # 📌 Ruta del archivo de stock crítico
ruta_pdf3 = os.path.join(ruta_carpeta,"Informe Stock Crítico - MP.pdf")
ruta_pdf4 = os.path.join(ruta_carpeta,"Informe Stock Crítico - PT.pdf")
ruta_pdf_bv = os.path.join(ruta_carpeta,"Informe Stock Crítico - BV.pdf")
ruta_pdf_compras = os.path.join(ruta_carpeta,"Informe Stock Crítico - COMPRAS.pdf")
ruta_dcrit = os.path.join(ruta_carpeta,"Destinatarios Critico.xlsx")
ruta_excel_pt = os.path.join(ruta_carpeta, "Informe PT.xlsx")
ruta_excel_bv = os.path.join(ruta_carpeta, "Informe BV.xlsx")
ruta_excel_compras = os.path.join(ruta_carpeta, "Informe Compras.xlsx")


### ----------------------Formato Excel--------------------

def aplicar_formato_excel(path_archivo, nombre_tabla="TablaDatos", estilo_tabla="TableStyleMedium9", aplicar_condicional=False):
    wb = load_workbook(path_archivo)

    for ws in wb.worksheets:  # 🔁 Recorre todas las hojas del archivo

        # Negrita encabezados
        for cell in ws[1]:
            cell.font = Font(bold=True)

        # Crear tabla
        max_row = ws.max_row
        max_col = ws.max_column
        ref = f"A1:{get_column_letter(max_col)}{max_row}"
        tabla = Table(displayName=f"{nombre_tabla}_{ws.title}", ref=ref)
        estilo = TableStyleInfo(name=estilo_tabla, showRowStripes=True)
        tabla.tableStyleInfo = estilo
        ws.add_table(tabla)

        # Congelar encabezados
        ws.freeze_panes = "A2"

        # Autoajuste de ancho de columnas
        for col in ws.columns:
            max_len = max(len(str(cell.value)) if cell.value else 0 for cell in col)
            ws.column_dimensions[get_column_letter(col[0].column)].width = max_len + 2

        # Opcional: formato condicional solo en RESUMEN
        if aplicar_condicional and ws.title == "RESUMEN":
            from openpyxl.formatting.rule import CellIsRule
            from openpyxl.styles import PatternFill
            rojo = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            ws.conditional_formatting.add(f"J2:J{max_row}", CellIsRule(operator='equal', formula=['"Crítico"'], fill=rojo))

    wb.save(path_archivo)
    print(f"✅ Formato aplicado a todas las hojas de: {path_archivo}")


##---------------------------------------------------------------


# ---------------------------CVS y EXCEL----------------------------------------

# 📌 Cargar historial anterior desde CSV si existe
if os.path.exists(ruta_csv):

    # 📌 PASO 1: Cargar historial desde el CVS (más rápido) y cambiar "," por ";"
    df_anterior = pd.read_csv(ruta_csv, sep=";") #Falta la concatenacion de la version anterior

    # Verificar si hay productos con problemas en la fecha.
    print("Fecha es NA en el df_anterior:")
    print(df_anterior[df_anterior["Fecha"].isna()])

    # 📌 PASO 2: Crear la columna Ultimo si no existe
    if "Ultimo" not in df_anterior.columns:
        df_anterior["Ultimo"] = False  # Inicializar la columna si no existe

    # 📌 PASO 3: Filtrar el historial solo para la última actualización.
    df_ultima_consulta = df_anterior[df_anterior["Ultimo"] == True] #Compara y crea una lista de verdaderos y falsos para ver cual es la ultima fecha.

    # 📌 PASO 4: Merge entre historial filtrado y stock actual para detectar cambios.
    # Hacer el merge solo con la última consulta y guardar todo, ambas partes con el merge "outer".
    df_comparado = df_ultima_consulta.merge(
        df_nuevo,
        on=["SKU", "UND", "Bodega"], # df_nuevo NO tiene la columna "Ultimo".
        how="outer",
        suffixes=("_anterior", "_nuevo")
        )
    
    # 🔧 PASO 5: Reconstruir columna Producto
    df_comparado["Producto"] = (
        df_comparado["Producto_nuevo"]
        .combine_first(df_comparado["Producto_anterior"])
    )

    # 🔧 PASO 6: Eliminar columnas duplicadas de Producto
    df_comparado.drop(
        columns=["Producto_anterior", "Producto_nuevo"],
        inplace=True,
        errors="ignore"
    )
        
    # 📌 PASO 7: Rellenar las filas con NA con 0.
    df_comparado["Stock_anterior"] = df_comparado["Stock_anterior"].fillna(0)
    df_comparado["Stock_nuevo"] = df_comparado["Stock_nuevo"].fillna(0)

    # 📌 PASO 8: Calcular ingreso/egreso
    df_comparado["Diferencia"] = df_comparado["Stock_nuevo"].fillna(0) - df_comparado["Stock_anterior"].fillna(0)

    # 📌 PASO 9: Crear la columna Movimientos si no existe
    if "Movimientos" not in df_comparado.columns:
        df_comparado["Movimientos"] = 0  # Inicializar la columna si no existe

    # 📌 PASO 10: Contador de movimientos (1 si hay cambio, 0 si no)
    df_comparado["Movimientos"] += df_comparado["Diferencia"].apply(lambda x: 1 if x != 0 else 0)

    # 📌 PASO 11: Actualizar la columna Ultimos, porque puede haber productos nuevos que lleguen y no tienen la columna ultimos.
    df_comparado["Ultimo"]=True

    # 📌 PASO 11: Eliminar columnas innecesarias
    df_comparado = df_comparado.drop(columns=["Fecha_anterior", "Stock_anterior"])

    # 📌 PASO 12: Renombrar columnas
    df_comparado = df_comparado.rename(columns={"Stock_nuevo": "Stock", "Fecha_nuevo": "Fecha"})

    # 📌 PASO 13: Filtrar el cuadro comparado para dejar solo los ultimos datos obtenidos y tratados. (quitar el merge)
    df_nuevos_movimientos = df_comparado[(df_comparado["Ultimo"] == True) & (df_comparado["Diferencia"] != 0)].copy() # Compara y crea una lista de verdaderos y falsos para ver cual es la ultima fecha.

    # 📌 PASO 14: Verificar si hay datos que guardar
    if not df_nuevos_movimientos.empty:

        # 📌 PASO 15: Crear una clave única por SKU y Bodega para comparar fácilmente
        df_anterior["clave"] = df_anterior["SKU"].astype(str) + "_" + df_anterior["Bodega"]
        df_nuevos_movimientos["clave"] = df_nuevos_movimientos["SKU"].astype(str) + "_" + df_nuevos_movimientos["Bodega"]

        # Solo poner "Ultimo = False" para los productos que cambiaron
        claves_actualizadas = set(df_nuevos_movimientos["clave"])
        df_anterior.loc[df_anterior["clave"].isin(claves_actualizadas), "Ultimo"] = False

        # (Opcional) eliminar la columna clave después
        df_anterior.drop(columns="clave", inplace=True)
        df_nuevos_movimientos.drop(columns="clave", inplace=True)

        # 📌 PASO 16: Guardar el historial actualizado
        df_total = pd.concat([df_anterior, df_nuevos_movimientos], ignore_index=True)

    else:
        print("📭 No hubo movimientos de stock en esta actualización.")
        df_total = df_anterior.copy()

else:

    # 📌 Si el historial no existe, aseguramos las columnas necesarias
    df_total = df_nuevo.copy()

    df_total["Diferencia"] = 0
    df_total["Movimientos"] = 0
    df_total["Ultimo"] = True

# 📌 PASO 17: Asegurar el orden correcto de las columnas (aplicado en ambos casos)
columnas_ordenadas = [
    "Fecha", 
    "Producto", 
    "SKU", 
    "UND", 
    "Bodega", 
    "Stock", 
    "Diferencia", 
    "Movimientos", 
    "Ultimo"]

df_total = df_total[[col for col in columnas_ordenadas if col in df_total.columns]]


#------------------------------------------------------------------
#--------------------PRODUCTOS INACTIVOS---------------------------
#------------------------------------------------------------------

# Convertir la columna "Fecha" a formato datetime, manejando errores con "coerce" para evitar problemas con valores no convertibles.
df_total["Fecha"] = pd.to_datetime(df_total["Fecha"], errors="coerce")

# Solo registros actuales
df_actual = df_total[df_total["Ultimo"] == True].copy()

# Fecha límite
fecha_corte = pd.Timestamp.now() - pd.DateOffset(months=6)

# Calcular días sin movimiento
df_actual["Dias_Sin_Movimiento"] = (
    pd.Timestamp.now() - df_actual["Fecha"]
).dt.days

# ---------------- INACTIVOS GENERALES -----------------------
df_inactivos = df_actual[df_actual["Fecha"] < fecha_corte]

# ---------------- INACTIVOS EN BODEGA VENTAS ----------------
df_inactivos_ventas = df_actual[
    (df_actual["Bodega"] == "Bodega Ventas") &
    (df_actual["Fecha"] < fecha_corte)
]


#------------------------------------------------------------------
#--------------------Hoja RESUMEN del EXCEL------------------------
#------------------------------------------------------------------

# Verificar si hay productos con problemas en la fecha.
print("Fecha es NA en el df_total antes del resumen:")
print(df_total[df_total["Fecha"].isna()][["Producto", "SKU", "UND", "Bodega"]])

# Filtrar por los registros marcados como "Ultimo"
df_total_ultimos = df_total[df_total["Ultimo"] == True]

# Verificar si hay productos con problemas en la fecha.
print("Fecha es NA en el df_total_ultimos después del filtro ultimos:")
print(df_total_ultimos[df_total_ultimos["Fecha"].isna()])

# Se salta los productos con problemas en la fecha.
df_total_ultimos["Fecha"] = pd.to_datetime(df_total_ultimos["Fecha"], errors="coerce")

# Para cada producto, SKU y UND, tomar la fecha más reciente.
fechas_recientes = df_total_ultimos.groupby(["Producto", "SKU", "UND"])["Fecha"].max().reset_index()

print(df_total_ultimos[df_total_ultimos["Fecha"].isna()])

# Tomar el último stock por producto y bodega
df_ultimos = df_total_ultimos.sort_values("Fecha").drop_duplicates(
    subset=["Producto", "SKU", "UND", "Bodega"], keep="last"
)

# Pivotear sin la fecha
df_resumen = df_ultimos.pivot_table(
    index=["Producto", "SKU", "UND"],
    columns="Bodega",
    values="Stock",
    aggfunc="sum",
    fill_value=0
).reset_index()

# Agregar fecha más reciente
df_resumen = df_resumen.merge(fechas_recientes, on=["Producto", "SKU", "UND"])

# Mover la columna de fecha al inicio
cols = ["Fecha"] + [col for col in df_resumen.columns if col != "Fecha"]
df_resumen = df_resumen[cols]

# Agregar una columna con el total de stock sumado
# 4: indica que se quieren todas las columnas desde la columna 4 en adelante (recordando que los índices empiezan en 0).
df_resumen["Stock"] = df_resumen.iloc[:, 4:].sum(axis=1)



# -----------------------VENTAS - ULTIMOS 3 MESES---------------------------

# 📌 Cargar Ventas si existe, y unirlo con "RESUMEN"
if os.path.exists(ruta_ventas):
    df_ventas = pd.read_excel(ruta_ventas, sheet_name="Sheet1", usecols="C,E")
    df_ventas.columns = ["SKU", "Cantidad_Venta"]  # Renombrar columnas para coincidir con df_resumen

    # Dividir Cantidad_Venta por 90 para obtener el promedio diario de ventas, y redondear a 2 decimales.
    df_ventas["Venta_Promedio_Diario"] = df_ventas["Cantidad_Venta"] / 90
    df_ventas["Venta_Promedio_Diario"] = df_ventas["Venta_Promedio_Diario"].round(2)

    # Eliminar la columna original de Cantidad_Venta, ya que no la necesitamos más.
    df_ventas.drop(columns=["Cantidad_Venta"], inplace=True)
    
    # 📌 Unir ventas con el RESUMEN usando "SKU".
    df_resumen = df_resumen.merge(
        df_ventas,
        on="SKU",
        how="left"
    )

# -----------------------CONSUMOS - ULTIMOS 3 MESES---------------------------

# 📌 Cargar Consumos si existe, y unirlo con "RESUMEN"
if os.path.exists(ruta_consumos):
    df_consumos = pd.read_excel(ruta_consumos, sheet_name="Consumos", usecols="H,K")
    df_consumos.columns = ["SKU", "Cantidad_Consumo"]  # Renombrar columnas para coincidir con df_resumen

    # Dividir Cantidad_Consumo por 90 para obtener el promedio diario de consumo, y redondear a 2 decimales.
    df_consumos["Consumo_Promedio_Diario"] = df_consumos["Cantidad_Consumo"] / 90
    df_consumos["Consumo_Promedio_Diario"] = df_consumos["Consumo_Promedio_Diario"].round(2)

    # Eliminar la columna original de Cantidad_Consumo, ya que no la necesitamos más.
    df_consumos.drop(columns=["Cantidad_Consumo"], inplace=True)
    
    # 📌 Unir consumos con el RESUMEN usando "SKU".
    df_resumen = df_resumen.merge(
        df_consumos,
        on="SKU",
        how="left"
    )

# -----------------------STOCK CRITICO---------------------------

# 📌 Cargar Stock Crítico si existe y unirlo con "RESUMEN"
if os.path.exists(ruta_stock_critico):
    df_stock_critico = pd.read_excel(
        ruta_stock_critico,
        sheet_name="Stock Critico",
        usecols="F,J,K,L,M,N,O,P,Q,R,U,Y"
        )
    df_stock_critico.columns = [
        "Estado",
        "Dias_Stock_Maximo", 
        "Dias_Stock_Minimo",
        "Tipo_Producto",
        "Stock_Minimo",
        "Stock_Maximo",
        "Produccion",
        "Compra",
        "Venta",
        "Informe", 
        "SKU", 
        "Estado_Variante"
        ]  # Renombrar columnas para coincidir con df_resumen

    # 📌 Unir stock crítico con el RESUMEN usando "SKU".
    df_resumen = df_resumen.merge(
        df_stock_critico,
        on="SKU",
        how="left")

    # 📌 Definir el denominador para evaluar si no da error.
    denominador1 = df_resumen["Stock_Minimo"]

    # Reemplazar ceros por NaN para evitar errores de división.
    denominador1 = denominador1.replace(0, np.nan)

    # 📌 -----------------FLUJO PROMEDIO DIARIO---------------

    # 📌 Definir el denominador para evaluar si no da error.
    df_resumen["Flujo_Promedio_Diario"] = df_resumen["Venta_Promedio_Diario"] + df_resumen["Consumo_Promedio_Diario"]

    # Reemplazar ceros por NaN para evitar errores de división.
    df_resumen["Flujo_Promedio_Diario"] = df_resumen["Flujo_Promedio_Diario"].replace(0, np.nan)

    # 📌 -----------------DÍAS DE STOCK ACTUAL---------------

    # 📌 Agregar columna Dias Stock Actual
    # Calcular el resultado
    df_resumen["Dias_Stock_Actual"] = df_resumen["Stock"] / df_resumen["Flujo_Promedio_Diario"]

    # 📌 ---------------Agregar columna %DISPONIBILIDAD-------------
    
    # Reemplazar ceros por NaN para evitar errores de división
    df_resumen["Stock_Maximo"] = df_resumen["Stock_Maximo"].replace(0, np.nan)

    # Calcular el resultado
    df_resumen["%Disponibilidad"] = (
        (df_resumen["Stock"] / df_resumen["Stock_Maximo"]) * 100
        ).round(1).apply(lambda x: f"{x}%" if not pd.isna(x) else "N/A")

    # ---------CANTIDAD COMPRAR O PRODUCIR--------
    # 📌 Agregar columna Cantidad_Comprar_o_Producir
    df_resumen["Cantidad_Comprar_o_Producir"] = (
        df_resumen["Stock_Maximo"] - df_resumen["Stock"]
        ).apply(lambda x: "LLENO" if x <= 0 else round(x, 2))

    # -------------------ALERTA STOCK----------------------

    # 📌 Crear columna de alerta: 1 si el stock total es menor al crítico, 0 si no
    df_resumen["Alerta_Stock"] = df_resumen.apply(
    lambda row: (
        "N/A"     if pd.isna(row["Stock_Minimo"])         else
        "QUIEBRE" if row["Stock"] == 0                    else
        "CRITICO" if row["Stock"] < row["Stock_Minimo"]   else
        "OK"      if row["Stock"] < row["Stock_Maximo"]*3 else
        "SOBRE"),
        axis=1
    )
    
# """     ("QUIEBRE" if row["Bodega Ventas"] == 0 else
#                 "CRITICO" if row["Bodega Ventas"] < row["Stock_Maximo_BV"] / 5 else
#                 "RELLENAR" if row["Bodega Ventas"] < row["Stock_Maximo_BV"] else
#                 "OK" if row["Bodega Ventas"] < row["Stock_Maximo_BV"] * 3 else
#                 "SOBRE"),
#                 axis=1 """


    # ---------------ENVIAR ALERTA-------------

    # 📌 Traer el estado anterior desde el excel creado
    if os.path.exists(ruta_excel):
        df_resumen_anterior = pd.read_excel(ruta_excel, sheet_name="RESUMEN")

        # 📌 Merge para hacer el match del critico anterior con el actual.
        df_comparado_critico = pd.merge(
            df_resumen_anterior, df_resumen,
            on="SKU",
            how="outer",
            suffixes=("_Anterior", ""))

        # 📌 Crea la lista de los productos que entraron de OK a Critico. Detectar cambio de OK ➡ CRITICO
        df_comparado_critico["Enviar_Alerta"] = (
            df_comparado_critico["Alerta_Stock_Anterior"] == "OK") & (df_comparado_critico["Alerta_Stock"] == "CRITICO")

        # 📌 Filtrar los productos a notificar
        df_a_notificar = df_comparado_critico[df_comparado_critico["Enviar_Alerta"] == True]

        # 📌 Filtrar solo las columnas necesarias para la notificación
        df_a_notificar = df_a_notificar[["Producto", "SKU", "Stock"]]


    # -----------------------------------------------------------
    # ----------------------BODEGA VENTAS-----------------------
    # -----------------------------------------------------------

    # 📌 ---------------DIAS STOCK ACTUAL BODEGA VENTAS----------

    # Definir el denominador para evaluar si no da error.
    denominador_dias_bv = df_resumen["Venta_Promedio_Diario"]

    # Reemplazar ceros por NaN para evitar errores de división
    denominador_dias_bv = denominador_dias_bv.replace(0, np.nan)

    # Calcular el resultado
    df_resumen["Dias_Stock_Actual_BV"] = df_resumen["Bodega Ventas"] / denominador_dias_bv


    # 📌 -----------------STOCK MAXIMO - BODEGA VENTAS---------------

    # Calcular STOCK MAXIMO BODEGA VENTAS
    df_resumen["Stock_Maximo_BV"] = df_resumen["Venta_Promedio_Diario"] *10

    # 📌 -----------------STOCK MINIMO - BODEGA VENTAS---------------

    # Calcular STOCK MINIMO BODEGA VENTAS
    df_resumen["Stock_Minimo_BV"] = df_resumen["Venta_Promedio_Diario"] * 2

    # 📌 -----------------STOCK SOBRE - BODEGA VENTAS---------------

    # Calcular STOCK SOBRE BODEGA VENTAS
    df_resumen["Stock_Sobre_BV"] = df_resumen["Stock_Maximo_BV"] * 3

    # 📌 ---------------%DISPONIBILIDAD - BODEGA VENTAS-------------

    # Reemplazar ceros por NaN para evitar errores de división
    denominador_smax_bv = df_resumen["Stock_Maximo_BV"].replace(0, np.nan)

    # Calcular el resultado
    df_resumen["%Disponibilidad_BV"] = (
        (df_resumen["Bodega Ventas"] / denominador_smax_bv) * 100
        ).round(1).apply(lambda x: f"{x}%" if not pd.isna(x) else "N/A")

    # ---------CANTIDAD SOLICITAR - BODEGA VENTAS--------
    # 📌 Agregar columna Cantidad_Solicitar
    df_resumen["Cantidad_Solicitar"] = (
        df_resumen["Stock_Maximo_BV"] - df_resumen["Bodega Ventas"]
        ).apply(lambda x: "LLENO" if x <= 0 else round(x, 2))
    
    # ------------ALERTA STOCK - BODEGA VENTAS---------
    # 📌 Crear columna de alerta: 1 si el stock total es menor al crítico, 0 si no
    df_resumen["Alerta_Stock_BV"] = df_resumen.apply(
    lambda row: (
        "QUIEBRE"  if row["Bodega Ventas"] == 0                     else
        "CRITICO"  if row["Bodega Ventas"] < row["Stock_Minimo_BV"] else
        "RELLENAR" if row["Bodega Ventas"] < row["Stock_Maximo_BV"] else
        "OK"       if row["Bodega Ventas"] < row["Stock_Sobre_BV"]  else
        "SOBRE"),
        axis=1
    )

    # 📌 En Stock_Maximo, reemplazamos NaN por "N/A"
    df_resumen["Stock_Maximo"] = df_resumen["Stock_Maximo"].fillna("N/A")

    # 📌 En Stock_Maximo_BV, reemplazamos NaN por "N/A"
    df_resumen["Stock_Maximo_BV"] = df_resumen["Stock_Maximo_BV"].fillna("N/A")

    # 📌 Asegurar el orden correcto de las columnas (aplicado en ambos casos)
    columnas_ordenadas = [
        "Fecha",
        "Tipo_Producto",
        "Producto",
        "SKU",
        "UND",
        "BODEGA PT",
        "Bodega General",
        "Bodega Producción",
        "Bodega Ventas",
        "Stock",
        "Venta_Promedio_Diario",
        "Dias_Stock_Minimo",
        "Stock_Minimo",
        "Dias_Stock_Actual",
        "Stock_Maximo",
        "Dias_Stock_Maximo",
        "Cantidad_Comprar_o_Producir",
        "%Disponibilidad",
        "Alerta_Stock",
        "Dias_Stock_Actual_BV",
        "Stock_Maximo_BV",
        "Cantidad_Solicitar",
        "%Disponibilidad_BV",
        "Alerta_Stock_BV",
        "Consumo_Promedio_Diario",
        "Estado",
        "Estado_Variante"]
    
    faltantes = [c for c in columnas_ordenadas if c not in df_resumen.columns]
    sobrantes = [c for c in df_resumen.columns if c not in columnas_ordenadas]

    if faltantes:
        raise ValueError(f"❌ Faltan columnas: {faltantes}")

    if sobrantes:
        print(f"⚠ Columnas extra detectadas: {sobrantes}")

    df_resumen = df_resumen[columnas_ordenadas]
    

else:
    print("⚠ No se encontró el archivo de stock crítico. Continuando sin él.")



#----------------------------------------------------------------------------------------------
# ------------------------------------GUARDAR LOS ARCHIVOS-------------------------------------
#----------------------------------------------------------------------------------------------

# -------------------------------------CVS y EXCEL---------------------------------------------

# 📌 Guardar en CSV, con delimitador “;”.
df_resumen[["Producto", "SKU", "UND"]].to_csv(ruta_productos_unicos, sep=";", index=False)

# 📌 Guardar en CSV, con delimitador “;”.
df_total.to_csv(ruta_csv, sep=";", index=False)

# 📌 Guardar en Excel con 2 hojas
with pd.ExcelWriter(ruta_excel, engine="openpyxl") as writer:
    df_total.to_excel(writer, sheet_name="BASE_DATOS", index=False)
    df_resumen.to_excel(writer, sheet_name="RESUMEN", index=False)

print(f"✅ Historial actualizado y guardado en:\n📂 {ruta_csv}\n📂 {ruta_excel}")

# -----------------FORMATO EXCEL - HISTORIAL---------------

aplicar_formato_excel(ruta_excel)





#-----------------------------------------------------------------------------
#----------------------------WORD Y PDF---------------------------------------
#-----------------------------------------------------------------------------

def repetir_encabezado_en_word(fila_encabezado):
    """
    Fuerza la primera fila de la tabla como encabezado repetido usando XML directo.
    """
    tr = fila_encabezado._tr  # Accede a la fila como XML <w:tr>
    trPr = tr.get_or_add_trPr()  # Obtiene o crea el nodo <w:trPr> (propiedades de la fila)
    tblHeader = OxmlElement('w:tblHeader')  # Crea el elemento <w:tblHeader>
    tblHeader.set(qn('w:val'), "true")  # Le asigna el atributo w:val="true"
    trPr.append(tblHeader)  # Agrega <w:tblHeader w:val="true"/> al nodo <w:trPr>





#-----------------------------------------------------------------------------
#---------------------WORD STOCK CRITICO - MP---------------------------------
#-----------------------------------------------------------------------------

### **🔹 Crear el Word**
doc = Document()


# -------------------CONFIGURACIÓN WORD STOCK CRITICO - MP-----------------------

# 1. Márgenes del documento, (1 cm = 0.3937 pulgadas)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# 2. Cambiar orientación a HORIZONTAL (LANDSCAPE)
section = doc.sections[-1]  # Puedes usar [-1] para tomar la última sección (o [0] si es una sola)
section.orientation = WD_ORIENT.LANDSCAPE

# ⚠️ Este paso es obligatorio para que el cambio surta efecto
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

# 3. Alineación del título principal
titulo = doc.add_heading("Informe Stock Crítico - MP", level=1)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar título

# (Opcional) Cambiar fuente y tamaño del título
run = titulo.runs[0]
#run.font.name = "Calibri"
run.font.size = Pt(16)


# -------------------CONFIGURACIÓN DATOS COLUMNAS - MP-----------------------

# Lista de encabezados de la hoja RESUMEN
df_encabezados3 = [
    "Producto",
    "SKU",
    "UND",
    "Stock",
    "Dias_Stock_Actual",
    "Stock_Maximo",
    "Cantidad_Comprar_o_Producir",
    "%Disponibilidad",
    "Alerta_Stock"]


# Lista de encabezados para el informe en word.
encabezados3_word = [
    "PRODUCTO",
    "SKU",
    "UND",
    "Stock",
    "Dias Stock",
    "Stock Max",
    "Comprar",
    "%Disp",
    "Estado"]

# Ordenar todos los datos de manera ascendente por DIAS STOCK ACTIAL
df_resumen_sorted3 = df_resumen.sort_values(by="Dias_Stock_Actual", ascending=True)  # Ascendente

# Filtrar los productos que tienen Stock Critico
df_resumen_sorted3_critico_mp = df_resumen_sorted3[
    (df_resumen_sorted3["Informe"] == "SI") &
    (df_resumen_sorted3["Tipo_Producto"] == "Bien") &
    (df_resumen_sorted3["Produccion"] == "MP") &
    (df_resumen_sorted3["Alerta_Stock"].isin(["QUIEBRE", "CRITICO"]))]

# Agregar tabla a Word
tabla = doc.add_table(rows=1, cols=len(encabezados3_word), style="Table Grid") #Tabla tamaño, 1 fila y el n° de encabezados por columnas

tabla.autofit = False  # ✅ Word NO ajustará automáticamente el ancho de cada columna

# Configurar ancho de las columnas
anchos_columnas_cm = [10, 2, 1.5, 2, 2, 2, 2, 2.12, 2.7]


# -------------------CONFIGURACIÓN ENCABEZADOS WORD - MP-----------------------

# ✅ Hace que el encabezado se repita en cada página
# tabla.rows[0].repeat_header = True  

# Agregar encabezados para el informe en word
celdas = tabla.rows[0].cells # Selecciona todas las celdas de la 1° fila (rows[0]) y de todas las columnas (cells).
for i, encabezado in enumerate(encabezados3_word):  # Recorre toda la lista de encabezados.
    celdas[i].text = str(encabezado) # Agrega cada encabezado de la lista a la respectiva celda.

    # ✅ Asignar ancho también a las celdas de las filas de datos
    celdas[i].width = Cm(anchos_columnas_cm[i])

    # Alineación vertical (celda)
    celdas[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  

    # 👉 Formato de fuente y alineación
    for paragraph in celdas[i].paragraphs:
        run = paragraph.runs[0]
        run.bold = True                      # Negrita
        #run.font.size = Pt(10)               # Tamaño de fuente
        #run.font.name = "Calibri"            # Tipo de letra
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrado


repetir_encabezado_en_word(tabla.rows[0])


# -------------------DATOS DE LA MATRIZ-----------------------

# Agregar filas de datos
for _, row in df_resumen_sorted3_critico_mp[df_encabezados3].iterrows():
    fila = tabla.add_row().cells
    for j, value in enumerate(row):
        
        # Si es la última columna (Estado/Alerta_Stock)
        if df_encabezados3[j] == "Alerta_Stock":
            if value == "CRITICO":
                fila[j].text = ""  # Dejamos el texto vacío inicialmente
                run = fila[j].paragraphs[0].add_run("❌ CRITICO")
                run.font.name = "Segoe UI Emoji"
            elif value == "OK":
                fila[j].text = ""
                run = fila[j].paragraphs[0].add_run("✅ OK")
                run.font.name = "Segoe UI Emoji"

        if pd.isna(value):                 # Si el valor en Nan entra al if y lo convierte en "N/A"
            fila[j].text = "N/A"
        elif isinstance(value, float):
            fila[j].text = f"{value:.1f}"  # Redondear a 1 decimal si es float
        else:
            fila[j].text = str(value)  # Dejar enteros y texto tal como están

        # ✅ Asignar ancho también a las celdas de las filas de datos
        fila[j].width = Cm(anchos_columnas_cm[j])

        # 👉 Estilo de cada celda
        for paragraph in fila[j].paragraphs:
            run = paragraph.runs[0]
            #run.bold = True                      # Negrita
            run.font.size = Pt(10)               # Tamaño de fuente
            #run.font.name = "Calibri"            # Tipo de letra
            
            # 👉 Alineación específica
            if j == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Primera columna a la izquierda
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Resto centrado


# --------------GUARDAR el Documento WORD STOCK CRITICO - MP----------------

doc.save(ruta_word3)
print(f"✅ Archivo guardado en: {ruta_word3}")


### -------------------------------PDF STOCK CRITICO - MP---------------------------

# Convierte un archivo Word específico a PDF
convert(ruta_word3, ruta_pdf3) # crea documento.pdf en el mismo directorio






#-----------------------------------------------------------------------------
#--------------------------WORD STOCK CRITICO - PT----------------------------
#-----------------------------------------------------------------------------

### **🔹 Crear el Word**
doc = Document()


# -------------------CONFIGURACIÓN WORD - PT-----------------------

# 1. Márgenes del documento, (1 cm = 0.3937 pulgadas)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# 2. Cambiar orientación a HORIZONTAL (LANDSCAPE)
section = doc.sections[-1]  # Puedes usar [-1] para tomar la última sección (o [0] si es una sola)
section.orientation = WD_ORIENT.LANDSCAPE

# ⚠️ Este paso es obligatorio para que el cambio surta efecto
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

# 3. Alineación del título principal
titulo = doc.add_heading("Informe Stock Crítico - PT", level=1)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar título

# (Opcional) Cambiar fuente y tamaño del título
run = titulo.runs[0]
run.font.name = "Calibri"
run.font.size = Pt(16)


# -------------------CONFIGURACIÓN DATOS COLUMNAS - PT-----------------------

# Lista de encabezados de la hoja RESUMEN
df_encabezados3 = [
    "Producto",
    "SKU",
    "UND",
    "Stock",
    "Dias_Stock_Actual",
    "Stock_Maximo",
    "Cantidad_Comprar_o_Producir",
    "%Disponibilidad",
    "Alerta_Stock"]


# Lista de encabezados para el informe en word.
encabezados3_word = [
    "PRODUCTO",
    "SKU",
    "UND",
    "Stock",
    "Dias Stock",
    "Stock Max",
    "Producir",
    "%Disp",
    "Estado"]

# Ordena todos los datos de manera ascendente por DIAS STOCK ACTIAL
df_resumen_sorted3 = df_resumen.sort_values(by="Dias_Stock_Actual", ascending=True)  # Ascendente

# Filtrar los productos que tienen Stock Critico
df_resumen_sorted3_critico_pt = df_resumen_sorted3[
                            (df_resumen_sorted3["Informe"] == "SI") &
                            (df_resumen_sorted3["Producto"].isin(["PT", "IM"]) ) &
                            (df_resumen_sorted3["Alerta_Stock"].isin(["QUIEBRE", "CRITICO"]))
                           ].sort_values(by="Dias_Stock_Actual")

# Agregar tabla a Word
tabla = doc.add_table(rows=1, cols=len(encabezados3_word), style="Table Grid") #Tabla tamaño, 1 fila y el n° de encabezados por columnas

tabla.autofit = False  # ✅ Word NO ajustará automáticamente el ancho de cada columna

anchos_columnas_cm = [10, 2, 1.5, 2, 2, 2, 2, 2.12, 2.7]


# ------------------- CONFIGURACIÓN ENCABEZADOS WORD - PT -----------------------

# Agregar encabezados para el informe en word
celdas = tabla.rows[0].cells # Selecciona todas las celdas de la 1° fila (rows[0]) y de todas las columnas (cells).
for i, encabezado in enumerate(encabezados3_word):  # Recorre toda la lista de encabezados.
    celdas[i].text = str(encabezado) # Agrega cada encabezado de la lista a la respectiva celda.

    # ✅ Asignar ancho también a las celdas de las filas de datos
    celdas[i].width = Cm(anchos_columnas_cm[i])

    # Alineación vertical (celda)
    celdas[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  

    # 👉 Formato de fuente y alineación
    for paragraph in celdas[i].paragraphs:
        run = paragraph.runs[0]
        run.bold = True                      # Negrita
        #run.font.size = Pt(10)               # Tamaño de fuente
        #run.font.name = "Calibri"            # Tipo de letra
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrado

repetir_encabezado_en_word(tabla.rows[0])

# -------------------DATOS DE LA MATRIZ - PT-----------------------

# Agregar filas de datos
for _, row in df_resumen_sorted3_critico_pt[df_encabezados3].iterrows():
    fila = tabla.add_row().cells
    for j, value in enumerate(row):

        # Si es la última columna (Estado/Alerta_Stock)
        if df_encabezados3[j] == "Alerta_Stock":
            if value == "CRITICO":
                fila[j].text = ""  # Dejamos el texto vacío inicialmente
                run = fila[j].paragraphs[0].add_run("❌ CRITICO")
                run.font.name = "Segoe UI Emoji"
            elif value == "OK":
                fila[j].text = ""
                run = fila[j].paragraphs[0].add_run("✅ OK")
                run.font.name = "Segoe UI Emoji"

        if pd.isna(value):                 # Si el valor en Nan entra al if y lo convierte en "N/A"
            fila[j].text = "N/A"
        elif isinstance(value, float):
            fila[j].text = f"{value:.1f}"  # Redondear a 1 decimal si es float
        else:
            fila[j].text = str(value)  # Dejar enteros y texto tal como están

        # ✅ Asignar ancho también a las celdas de las filas de datos
        fila[j].width = Cm(anchos_columnas_cm[j])

        # 👉 Estilo de cada celda
        for paragraph in fila[j].paragraphs:
            run = paragraph.runs[0]
            #run.bold = True                      # Negrita
            run.font.size = Pt(10)               # Tamaño de fuente
            #run.font.name = "Calibri"            # Tipo de letra
            
            # 👉 Alineación específica
            if j == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Primera columna a la izquierda
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Resto centrado


# --------------GUARDAR el Documento WORD STOCK CRITICO - PT----------------

doc.save(ruta_word4)
print(f"✅ Archivo guardado en: {ruta_word4}")


### -------------------------------PDF STOCK CRITICO - PT---------------------------

# Convierte un archivo Word específico a PDF
convert(ruta_word4, ruta_pdf4) # crea documento.pdf en el mismo directorio


##--------------------------------------------------------------------------
##------------------------EXCEL STOCK CRITICO - PT--------------------------
##--------------------------------------------------------------------------

# # Asegúrate de que este sea el mismo orden de columnas en tu DataFrame
# columnas_mapeo = {
#     "PRODUCTO": "Producto",
#     "SKU": "SKU",
#     "UND": "UND",
#     "Stock": "Bodega Ventas",
#     "Dias Stock": "Dias_Stock_Actual",
#     "Stock Max": "Stock_Maximo",
#     "Producir": "Cantidad_Comprar_o_Producir",
#     "%Disp": "%Disponibilidad",
#     "Estado MarJano": "Alerta_Stock"
# }

# # Ordenar todos los datos de manera ascendente por DIAS STOCK ACTIAL
# df_resumen_pt_ordenado = df_resumen.sort_values(by="Dias_Stock_Actual", ascending=True)  # Ascendente

# # Filtrar los productos que tienen Stock Critico
# df_resumen_pt_ordenado_filtrado = df_resumen_pt_ordenado[
#                             (df_resumen_pt_ordenado["Tipo"].isin(["PT", "IM"]) ) &
#                             (df_resumen_pt_ordenado["Alerta_Stock"].isin(["CRITICO"]))]

# # Crear DataFrame para exportar: seleccionar columnas y renombrarlas
# df_exportar = df_resumen_pt_ordenado_filtrado[list(columnas_mapeo.values())].copy()
# df_exportar.columns = list(columnas_mapeo.keys())

# # Exportar a Excel
# df_exportar.to_excel(ruta_excel_pt, index=False)

# print(f"✅ Informe exportado exitosamente a: {ruta_excel}")

# aplicar_formato_excel(ruta_excel_pt)





#-----------------------------------------------------------------------------
#-------------------WORD STOCK CRITICO - BODEGA VENTAS------------------------
#-----------------------------------------------------------------------------

### **🔹 Crear el Word**
doc = Document()


# ---------------CONFIGURACIÓN WORD STOCK CRITICO - BODEGA VENTAS------------------

# 1. Márgenes del documento, (1 cm = 0.3937 pulgadas)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# 2. Cambiar orientación a HORIZONTAL (LANDSCAPE)
section = doc.sections[-1]  # Puedes usar [-1] para tomar la última sección (o [0] si es una sola)
section.orientation = WD_ORIENT.LANDSCAPE

# ⚠️ Este paso es obligatorio para que el cambio surta efecto
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

# 3. Alineación del título principal
titulo = doc.add_heading("Informe Stock Crítico - Bodega Ventas (BV)", level=1)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar título

# (Opcional) Cambiar fuente y tamaño del título
run = titulo.runs[0]
#run.font.name = "Calibri"
run.font.size = Pt(16)


# -------------------CONFIGURACIÓN DATOS COLUMNAS - BODEGA VENTAS-----------------------

# Lista de encabezados de la hoja RESUMEN
df_encabezados_bv = [
    "Producto",
    "SKU",
    "UND",
    "Bodega Ventas",
    "Dias_Stock_Actual_BV",
    "Stock_Maximo_BV", 
    "Cantidad_Solicitar",
    "%Disponibilidad_BV",
    "Alerta_Stock_BV"]


# Lista de encabezados para el informe en word.
encabezados_word_bv = [
    "PRODUCTO",
    "SKU",
    "UND",
    "Stock BV", 
    "Dias Stock",
    "Stock Max", 
    "Cant. Solicitar",
    "%Disp",
    "Estado MarJano",
    ""]

# Ordenar todos los datos de manera ascendente por DIAS STOCK ACTIAL
df_resumen_ordenado_bv = df_resumen.sort_values(by="Dias_Stock_Actual_BV", ascending=True)  # Ascendente

# Filtrar los productos que tienen Stock Critico
df_resumen_ordenado_filtrado_bv = df_resumen_ordenado_bv[
                            (df_resumen_ordenado_bv["Informe"] == "SI") &
                            (df_resumen_ordenado_bv["Venta_Promedio_Diario"] > 0)
                            ]

# Agregar tabla a Word
tabla = doc.add_table(rows=1, cols=len(encabezados_word_bv), style="Table Grid") #Tabla tamaño, 1 fila y el n° de encabezados por columnas

tabla.autofit = False  # ✅ Word NO ajustará automáticamente el ancho de cada columna

# Configurar ancho de las columnass
anchos_columnas_cm = [7, 1.69, 1.48, 1.82, 2, 2, 2, 1.94, 2.4, 3.99]


# -------------------CONFIGURACIÓN ENCABEZADOS WORD - BODEGA VENTAS-----------------------

# Agregar encabezados para el informe en word
celdas = tabla.rows[0].cells # Selecciona todas las celdas de la 1° fila (rows[0]) y de todas las columnas (cells).
for i, encabezado in enumerate(encabezados_word_bv):  # Recorre toda la lista de encabezados.
    celdas[i].text = str(encabezado) # Agrega cada encabezado de la lista a la respectiva celda.

    # ✅ Asignar ancho también a las celdas de las filas de datos
    celdas[i].width = Cm(anchos_columnas_cm[i])

    # Alineación vertical (celda)
    celdas[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  

    # 👉 Formato de fuente y alineación
    for paragraph in celdas[i].paragraphs:
        run = paragraph.runs[0]
        run.bold = True                      # Negrita
        #run.font.size = Pt(10)               # Tamaño de fuente
        #run.font.name = "Calibri"            # Tipo de letra
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrado


repetir_encabezado_en_word(tabla.rows[0])


# -------------------DATOS DE LA MATRIZ - BODEGA VENTAS-----------------------

# Agregar filas de datos
for _, row in df_resumen_ordenado_filtrado_bv[df_encabezados_bv].iterrows():
    fila = tabla.add_row().cells
    for j, value in enumerate(row):
        
        # Si es la última columna (Estado/Alerta_Stock)
        if df_encabezados_bv[j] == "Alerta_Stock":
            if value == "CRITICO":
                fila[j].text = ""  # Dejamos el texto vacío inicialmente
                run = fila[j].paragraphs[0].add_run("❌ CRITICO")
                run.font.name = "Segoe UI Emoji"
            elif value == "OK":
                fila[j].text = ""
                run = fila[j].paragraphs[0].add_run("✅ OK")
                run.font.name = "Segoe UI Emoji"

        if pd.isna(value):                 # Si el valor en Nan entra al if y lo convierte en "N/A"
            fila[j].text = "N/A"
        elif isinstance(value, float):
            fila[j].text = f"{value:.1f}"  # Redondear a 1 decimal si es float
        else:
            fila[j].text = str(value)  # Dejar enteros y texto tal como están

        # ✅ Asignar ancho también a las celdas de las filas de datos
        fila[j].width = Cm(anchos_columnas_cm[j])

        # 👉 Estilo de cada celda
        for paragraph in fila[j].paragraphs:
            run = paragraph.runs[0]
            #run.bold = True                      # Negrita
            run.font.size = Pt(10)               # Tamaño de fuente
            #run.font.name = "Calibri"            # Tipo de letra
            
            # 👉 Alineación específica
            if j == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Primera columna a la izquierda
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Resto centrado


# --------------GUARDAR el Documento WORD STOCK CRITICO - BODEGA VENTAS----------------

doc.save(ruta_word_bv)
print(f"✅ Archivo guardado en: {ruta_word_bv}")


### -------------------------------PDF STOCK CRITICO - BODEGA VENTAS---------------------------

# Convierte un archivo Word específico a PDF
convert(ruta_word_bv, ruta_pdf_bv) # crea documento.pdf en el mismo directorio


### -----------------------------EXCEL STOCK CRITICO - BODEGA VENTAS-----------------------------

# Asegúrate de que este sea el mismo orden de columnas en tu DataFrame
columnas_mapeo = {
    "PRODUCTO": "Producto",
    "SKU": "SKU",
    "UND": "UND",
    "Stock BV": "Bodega Ventas",
    "Dias Stock": "Dias_Stock_Actual_BV",
    "Stock Max": "Stock_Maximo_BV",
    "Cant. solicitar": "Cantidad_Solicitar",
    "%Disp": "%Disponibilidad_BV",
    "Estado": "Alerta_Stock_BV"
}

# Ordenar todos los datos de manera ascendente por DIAS STOCK ACTIAL
df_resumen_ordenado_bv = df_resumen.sort_values(by="Dias_Stock_Actual_BV", ascending=True)  # Ascendente

# Filtrar los productos que tienen Stock Critico
df_resumen_ordenado_filtrado_bv = df_resumen_ordenado_bv[
                            (df_resumen_ordenado_bv["Venta_Promedio_Diario"] > 0)]


# Crear DataFrame para exportar: seleccionar columnas y renombrarlas
df_exportar = df_resumen_ordenado_filtrado_bv[list(columnas_mapeo.values())].copy()
df_exportar.columns = list(columnas_mapeo.keys())

# Exportar a Excel
df_exportar.to_excel(ruta_excel_bv, index=False)

print(f"✅ Informe exportado exitosamente a: {ruta_excel_bv}")

aplicar_formato_excel(ruta_excel_bv)





#-------------------------------------------------------------------------------------------------
#------------------------------WORD STOCK CRITICO - COMPRAS---------------------------------------
#-------------------------------------------------------------------------------------------------

### **🔹 Crear el Word**
doc = Document()


# -------------------CONFIGURACIÓN WORD STOCK CRITICO - COMPRAS-----------------------

# 1. Márgenes del documento, (1 cm = 0.3937 pulgadas)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# 2. Cambiar orientación a HORIZONTAL (LANDSCAPE)
section = doc.sections[-1]  # Puedes usar [-1] para tomar la última sección (o [0] si es una sola)
section.orientation = WD_ORIENT.LANDSCAPE

# ⚠️ Este paso es obligatorio para que el cambio surta efecto
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

# 3. Alineación del título principal
titulo = doc.add_heading("Informe Stock Crítico - COMPRAS", level=1)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar título

# (Opcional) Cambiar fuente y tamaño del título
run = titulo.runs[0]
#run.font.name = "Calibri"
run.font.size = Pt(16)


# -------------------CONFIGURACIÓN DATOS COLUMNAS - COMPRAS-----------------------

# Lista de encabezados de la hoja RESUMEN
df_encabezados_compras = [
    "Producto",
    "SKU",
    "UND",
    "Stock",
    "Dias_Stock_Actual",
    "Stock_Maximo", 
    "Cantidad_Comprar_o_Producir",
    "%Disponibilidad",
    "Alerta_Stock"]


# Lista de encabezados para el informe en word.
encabezados_compras_word = [
    "PRODUCTO",
    "SKU",
    "UND",
    "Stock",
    "Dias Stock",
    "Stock Max",
    "Comprar",
    "%Disp",
    "Estado"]

# Ordenar todos los datos de manera ascendente por DIAS STOCK ACTIAL
df_resumen_compras_ordenado = df_resumen.sort_values(by="Dias_Stock_Actual", ascending=True)  # Ascendente

# Filtrar los productos que tienen Stock Critico
df_resumen_compras_ordenado_filtrado = df_resumen_compras_ordenado[
    (df_resumen_compras_ordenado["Informe"] == "SI") &
    (df_resumen_compras_ordenado["Tipo_Producto"].isin(["Bien"])) &
    (df_resumen_compras_ordenado["Compra"].isin(["SI"])) &
    (df_resumen_compras_ordenado["Alerta_Stock"].isin(["QUIEBRE", "CRITICO"]))]

# Agregar tabla a Word
tabla = doc.add_table(rows=1, cols=len(encabezados_compras_word), style="Table Grid") #Tabla tamaño, 1 fila y el n° de encabezados por columnas

tabla.autofit = False  # ✅ Word NO ajustará automáticamente el ancho de cada columna

# Configurar ancho de las columnas
anchos_columnas_cm = [10, 2, 1.5, 2, 2, 2, 2, 2.12, 2.7]


# -------------------CONFIGURACIÓN ENCABEZADOS WORD - COMPRAS-----------------------

# ✅ Hace que el encabezado se repita en cada página
# tabla.rows[0].repeat_header = True

# Agregar encabezados para el informe en word
celdas = tabla.rows[0].cells # Selecciona todas las celdas de la 1° fila (rows[0]) y de todas las columnas (cells).
for i, encabezado in enumerate(encabezados_compras_word):  # Recorre toda la lista de encabezados.
    celdas[i].text = str(encabezado) # Agrega cada encabezado de la lista a la respectiva celda.

    # ✅ Asignar ancho también a las celdas de las filas de datos
    celdas[i].width = Cm(anchos_columnas_cm[i])

    # Alineación vertical (celda)
    celdas[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER  

    # 👉 Formato de fuente y alineación
    for paragraph in celdas[i].paragraphs:
        run = paragraph.runs[0]
        run.bold = True                      # Negrita
        #run.font.size = Pt(10)               # Tamaño de fuente
        #run.font.name = "Calibri"            # Tipo de letra
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrado


repetir_encabezado_en_word(tabla.rows[0])


# -------------------DATOS DE LA MATRIZ - COMPRAS-----------------------

# Agregar filas de datos
for _, row in df_resumen_compras_ordenado_filtrado[df_encabezados_compras].iterrows():
    fila = tabla.add_row().cells
    for j, value in enumerate(row):
        
        # Si es la última columna (Estado/Alerta_Stock)
        if df_encabezados_compras[j] == "Alerta_Stock":
            if value == "CRITICO":
                fila[j].text = ""  # Dejamos el texto vacío inicialmente
                run = fila[j].paragraphs[0].add_run("❌ CRITICO")
                run.font.name = "Segoe UI Emoji"
            elif value == "OK":
                fila[j].text = ""
                run = fila[j].paragraphs[0].add_run("✅ OK")
                run.font.name = "Segoe UI Emoji"

        if pd.isna(value):                 # Si el valor en Nan entra al if y lo convierte en "N/A"
            fila[j].text = "N/A"
        elif isinstance(value, float):
            fila[j].text = f"{value:.1f}"  # Redondear a 1 decimal si es float
        else:
            fila[j].text = str(value)  # Dejar enteros y texto tal como están

        # ✅ Asignar ancho también a las celdas de las filas de datos
        fila[j].width = Cm(anchos_columnas_cm[j])

        # 👉 Estilo de cada celda
        for paragraph in fila[j].paragraphs:
            run = paragraph.runs[0]
            #run.bold = True                      # Negrita
            run.font.size = Pt(10)               # Tamaño de fuente
            #run.font.name = "Calibri"            # Tipo de letra
            
            # 👉 Alineación específica
            if j == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Primera columna a la izquierda
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Resto centrado


# --------------GUARDAR el Documento WORD STOCK CRITICO - COMPRAS----------------

doc.save(ruta_word_compras)
print(f"✅ Archivo guardado en: {ruta_word_compras}")


### -------------------------------PDF STOCK CRITICO - COMPRAS---------------------------

# Convierte un archivo Word específico a PDF
convert(ruta_word_compras, ruta_pdf_compras) # crea documento.pdf en el mismo directorio



#--------------------------------------------------------------------------------------------
# ### --------------------------EXCEL STOCK CRITICO - COMPRAS--------------------------------
#--------------------------------------------------------------------------------------------

# # Asegúrate de que este sea el mismo orden de columnas en tu DataFrame
# columnas_mapeo = {
#     "PRODUCTO": "Producto",
#     "SKU": "SKU",
#     "UND": "UND",
#     "Stock MarJano": "Stock",
#     "Dias Stock": "Dias_Stock_Actual",
#     "Stock Max": "Stock_Maximo",
#     "Comprar": "Cantidad_Comprar_o_Producir",
#     "%Disp": "%Disponibilidad",
#     "Estado": "Alerta_Stock"
# }

# # Ordenar todos los datos de manera ascendente por DIAS STOCK ACTIAL
# df_resumen_compras_ordenado = df_resumen.sort_values(by="Dias_Stock_Actual", ascending=True)  # Ascendente

# # Filtrar los productos que tienen Stock Critico
# df_resumen_compras_ordenado_filtrado = df_resumen_compras_ordenado[
#                             (df_resumen_compras_ordenado["Solicitud"].isin(["COMPRA", "AMBOS"])) &
#                             (df_resumen_compras_ordenado["Alerta_Stock"].isin(["CRITICO"]))]


# # Crear DataFrame para exportar: seleccionar columnas y renombrarlas
# df_exportar = df_resumen_compras_ordenado_filtrado[list(columnas_mapeo.values())].copy()
# df_exportar.columns = list(columnas_mapeo.keys())


# # Exportar a Excel
# df_exportar.to_excel(ruta_excel_compras, index=False)

# print(f"✅ Informe exportado exitosamente a: {ruta_excel_compras}")

# aplicar_formato_excel(ruta_excel_compras)









#----------------------------------------------------------
### -----------------------EMAIL--------------------------


# -----------CONFIGURACIÓN----------

# 📌 Configurar los datos del correo
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
EMAIL_SENDER = "tomas.guevara.mesa@gmail.com"
EMAIL_PASSWORD = "ydwd ntns xjia zscu"  # Contraseña de aplicación (secreta)

# 📌 Leer destinatarios
df_destinatarios = pd.read_excel(ruta_dcrit)


# 📌 Iniciar servidor SMTP
try:
    server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
    server.starttls()
    server.login(EMAIL_SENDER, EMAIL_PASSWORD)



    ## -------------ALERTA STOCK CRITICO------------

    # Generar tabla HTML de alerta si hay productos críticos
    tabla_html = ""
    if not df_a_notificar.empty:
        raw_html = df_a_notificar.to_html(index=False, border=0, justify="center")
        tabla_html = raw_html.replace("<td>", '<td style="padding: 6px 12px;">') \
                             .replace("<th>", '<th style="padding: 6px 12px; background-color:#f0f0f0;">')


    ## --------------INFORMES STOCK CRITICO-------------

    for _, row in df_destinatarios.iterrows():
            nombre = row["Nombre"]
            email = row["Email"]
            informe_mp = row["SC - MP"].strip().upper() == "SI"
            informe_pt = row["SC - PT"].strip().upper() == "SI"
            informe_bv = row["SC - BV"].strip().upper() == "SI"
            informe_compras = row["SC - COMPRAS"].strip().upper() == "SI"
            excel_pt = row["EXCEL - PT"].strip().upper() == "SI"
            excel_compras = row["EXCEL - COMPRAS"].strip().upper() == "SI"
            excel_historial = row["EXCEL - HISTORIAL"].strip().upper() == "SI"
            excel_bv = row["EXCEL - BV"].strip().upper() == "SI"
            alerta_stock = row["Alerta Stock"].strip().upper() == "SI"

            # Solo si le corresponde recibir alerta y hay productos críticos
            alerta_html = ""
            if alerta_stock and tabla_html:
                alerta_html = f"""
                <p>🚨 Te informamos que los siguientes productos han cambiado su estado de <b>OK</b> a <b>CRÍTICO</b>:</p>
                {tabla_html}
                <p>Por favor, revisa esta situación para tomar las medidas correspondientes.</p>
                """
                
            # Construir mensaje HTML completo
            html_message = f"""
            <html>
            <body>
                <p>Hola {nombre},</p>
                {alerta_html}
                <p>Te comparto los informes actualizados relacionados con el stock de productos:</p>
                <ul>
                    {"<li><b>Informe Stock Crítico - MP:</b> incluye los productos (en orden por días) que requieren atención inmediata.</li>" if informe_mp else ""}
                    {"<li><b>Informe Stock Crítico - PT:</b> incluye los productos (en orden por días) que requieren atención inmediata.</li>" if informe_pt else ""}
                    {"<li><b>Informe Stock Crítico - BV:</b> incluye los productos (en orden por días) que requieren atención inmediata.</li>" if informe_bv else ""}
                    {"<li><b>Informe Stock Crítico - COMPRAS:</b> incluye los productos (en orden por días) que requieren atención inmediata.</li>" if informe_compras else ""}
                    {"<li><b>Excel - PT:</b> excel del Informe Stock Crítico - PT.</li>" if excel_pt else ""}
                    {"<li><b>Excel - COMPRAS:</b> excel del Informe Stock Crítico - COMPRAS.</li>" if excel_compras else ""}
                    {"<li><b>Excel - BODEGA VENTAS:</b> excel del Informe Stock Crítico - BODEGA VENTAS.</li>" if excel_bv else ""}
                    {"<li><b>Historial Completo:</b> muestra el seguimiento general del inventario.</li>" if excel_historial else ""}
                </ul>
                <p>Por favor, revisa los documentos adjuntos. Ante cualquier consulta, no dudes en contactarme.</p>
                <p>Saludos cordiales,<br><b>Tomás Guevara</b></p>
            </body>
            </html>
            """

            msg = EmailMessage()
            msg["Subject"] = f"📦 Informe Stock Crítico"
            msg["From"] = EMAIL_SENDER
            msg["To"] = email
            msg.set_content("Este mensaje está en formato HTML.")
            msg.add_alternative(html_message, subtype="html")

            # Adjuntar solo lo que corresponde
            if informe_mp and os.path.exists(ruta_pdf3):
                with open(ruta_pdf3, "rb") as f:
                    msg.add_attachment(f.read(), 
                                       maintype="application", 
                                       subtype="pdf", 
                                       filename="Informe MP.pdf")

            if informe_pt and os.path.exists(ruta_pdf4):
                 with open(ruta_pdf4, "rb") as f:
                    msg.add_attachment(f.read(), 
                                       maintype="application", 
                                       subtype="pdf", 
                                       filename="Informe PT.pdf")
                    
            if informe_bv and os.path.exists(ruta_pdf_bv):
                 with open(ruta_pdf_bv, "rb") as f:
                    msg.add_attachment(f.read(), 
                                       maintype="application", 
                                       subtype="pdf", 
                                       filename="Informe BV.pdf")
                    
            if informe_compras and os.path.exists(ruta_pdf_compras):
                 with open(ruta_pdf_compras, "rb") as f:
                    msg.add_attachment(f.read(), 
                                       maintype="application", 
                                       subtype="pdf", 
                                       filename="Informe COMPRAS.pdf")

            # Adjuntar Excel del historial
            if excel_historial:
                if os.path.exists(ruta_excel):
                    with open(ruta_excel, "rb") as f:
                        msg.add_attachment(
                            f.read(),
                            maintype="application",
                            subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            filename="Historial Stock Critico.xlsx"
                        )
                else:
                    print("⚠️ El archivo del historial no fue encontrado.")

            # Adjuntar Excel de Informe PT
            if excel_pt:
                if os.path.exists(ruta_excel_pt):
                    with open(ruta_excel_pt, "rb") as f:
                        msg.add_attachment(
                            f.read(),
                            maintype="application",
                            subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            filename="Informe PT.xlsx"
                        )
                else:
                    print("⚠️ El archivo Informe PT no fue encontrado.")

            # Adjuntar Excel de Informe COMPRAS
            if excel_compras:
                if os.path.exists(ruta_excel_compras):
                    with open(ruta_excel_compras, "rb") as f:
                        msg.add_attachment(
                            f.read(),
                            maintype="application",
                            subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            filename="Informe COMPRAS.xlsx"
                        )
                else:
                    print("⚠️ El archivo Excel COMPRAS no fue encontrado.")
            
            # Adjuntar Excel de Informe BODEGA VENTAS
            if excel_bv:
                if os.path.exists(ruta_excel_bv):
                    with open(ruta_excel_bv, "rb") as f:
                        msg.add_attachment(
                            f.read(),
                            maintype="application",
                            subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            filename="Informe BODEGA VENTAS.xlsx"
                        )
                else:
                    print("⚠️ El archivo Excel BODEGA VENTAS no fue encontrado.")

            server.send_message(msg)
            print(f"✅ Correo enviado a {nombre} ({email})")
    

    server.quit()
    print("📧 Todos los correos fueron enviados correctamente.")

except Exception as e:
    print("❌ Error al enviar los correos:", str(e))


######## -----------------------------Bucle-----------------------------------------

# # 🔄 **Bucle infinito con opción de salida**
# try:
#     while True:
#         actualizar_stock()
#         print(f"⏳ Esperando {INTERVALO_SEGUNDOS} segundos antes de la próxima actualización...")
#         print("📢 Escribe 'salir' y presiona Enter para detener el programa.")
        
#         # Esperar el intervalo definido y revisar si el usuario quiere salir
#         for _ in range(INTERVALO_SEGUNDOS):
#             time.sleep(1)  
#             if os.path.exists("salir.txt"):  # Detectar si se crea un archivo de salida
#                 os.remove("salir.txt")
#                 raise KeyboardInterrupt  # Detener el programa con una interrupción

# except KeyboardInterrupt:
#     print("\n⏹ Programa detenido manualmente. Finalizando ejecución...")